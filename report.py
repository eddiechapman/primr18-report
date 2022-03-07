#!/usr/bin/env python3

import argparse
import csv
from dataclasses import dataclass
import pathlib
from pprint import pprint

import docx
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn


column_names = {
    "case_1": {"irb_consideration": "Q137", "key_factors": "Q138_13_TEXT", "ethical_concerns": "Q141"},
    "case_2": {"irb_consideration": "20", "key_factors": "21_13_TEXT", "ethical_concerns": "Q115"},
    "case_3": {"irb_consideration": "40", "key_factors": "41_13_TEXT", "ethical_concerns": "Q125"},
    "case_4": {"irb_consideration": "30", "key_factors": "31_13_TEXT", "ethical_concerns": "Q121"},
    "case_5": {"irb_consideration": "45", "key_factors": "46_13_TEXT", "ethical_concerns": "Q127"},
    "case_6": {"irb_consideration": "50", "key_factors": "51_13_TEXT", "ethical_concerns": "Q129"},
    "case_7": {"irb_consideration": "65", "key_factors": "66_13_TEXT", "ethical_concerns": "Q135"},
    "case_8": {"irb_consideration": "55", "key_factors": "56_13_TEXT", "ethical_concerns": "Q131"},
    "case_9": {"irb_consideration": "25", "key_factors": "26_13_TEXT", "ethical_concerns": "Q117"},
    "case_10": {"irb_consideration": "35", "key_factors": "36_13_TEXT", "ethical_concerns": "Q123"},
    "case_11": {"irb_consideration": "60", "key_factors": "61_13_TEXT", "ethical_concerns": "Q133"}
}

case_study_info = {
    "case_1": {"title": "Predict election via news comments", "prompt": "Researchers plan to scrape public comments from online newspaper pages to predict election outcomes. They will aggregate their analysis to determine public sentiment. The researchers don’t plan to inform commenters, and they plan to collect potentially-identifiable user names. Scraping comments violates the newspaper’s terms of service."},
    "case_2": {"title": "Predict risky drug-use via Twitter", "prompt": "Researchers plan to scrape public Twitter feeds to predict risky drug-use behaviors. They will analyze individual behaviors. The researchers don’t plan to inform Twitter users, but they will not collect any identifying information. Scraping Tweets does not violate Twitter’s terms of service."},
    "case_3": {"title": "Study sexual behavior via dating app data", "prompt": "Researchers plan to analyze private interaction data from a dating site to understand the sexual behavior of groups. The researchers plan to collect informed consent from dating site users, and they plan to collect identifiable information from participants. Asking users for permission to use their data does not violate the dating site’s terms of service."},
    "case_4": {"title": "Understand political views via news comments", "prompt": "Researchers plan to collect newspaper comments by reading articles and cutting and pasting all associated comments into spreadsheets. They will use qualitative analysis to understand individual political views. The researchers don’t plan to inform commenters, and they plan to collect potentially-identifiable user names. Cutting and pasting comments does not violate the newspaper’s terms of service."},
    "case_5": {"title": "Study group mobility via cell phone geolocation data", "prompt": "Researchers plan to work with a mobile phone company to collect geolocation data to understand group mobility patterns in a city. The researchers will not inform the mobile phone users, and they will not collect any additional identifying information. Partnering with the mobile phone company to collect data does not violate the company’s terms of service."},
    "case_6": {"title": "Predict student mental health via health records and social media", "prompt": "Researchers plan to combine mental health records provided by a university and public social media activity to predict mental health conditions among students. The researchers plan to collect informed consent, and they plan to collect identifiable information from participants."},
    "case_7": {"title": "Study political event via public Tweets", "prompt": "Researchers plan to use a database of public tweets curated and shared by another researcher to study a political event. Researchers do not plan to inform the original posters, and researchers have taken measures to de-identify the data."},
    "case_8": {"title": "Predict mental health via health forum data and public Tweets", "prompt": "Researchers plan to scrape data from an open health forum and combine it with scraped tweets to predict mental health conditions. The researchers will not inform forum users, and they may collect potentially identifying information. Scraping data violates neither the health forum nor Twitter’s terms of service."},
    "case_9": {"title": "Predict sexual preference via profile photos", "prompt": "Researchers plan to scrape profile photos, which are visible to any member of the service, from a dating site to build models that predict sexual preference or behavior. Researchers will not inform the dating site users, but they will not collect any identifying information and their photograph dataset will not be released publicly. Creating a fake profile, necessary to access the photos, violates the dating site’s terms of service."},
    "case_10": {"title": "Study impact of exercise via Apple Healthkit data", "prompt": "Researchers plan to ask Apple HealthKit users to voluntarily submit their activity data to understand the general impact of exercise on a health condition. The researchers plan to obtain informed consent, and they plan to collect identifiable information from participants. Asking users to submit activity data does not violate Apple Health Kit’s terms of service."},
    "case_11": {"title": "Study group dynamics via Facebook posts", "prompt": "Researchers plan to scrape public posts and interactions from Facebook to study group-level dynamics. They plan to collect informed consent from the original poster, but not those they interacted with, and they may collect identifying information. Scraping posts with permission of the original poster does not violate Facebook’s terms of service."}
}


def add_horizontal_line(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)


@dataclass
class CaseStudyResponse:
    respondent_id: str
    irb_consideration: str
    key_factors: str
    ethical_concerns: str


def main(args):
    responses = {f"case_{i}": [] for i in range(1, 12)}

    with args.infile.open("r") as f:
        reader = csv.DictReader(f)
        next(reader)
        next(reader)
        
        for row in reader:
            for i in range(1, 12):
                case_study = f"case_{i}"
                respondent = row["ResponseId"]
                irb_consideration = row[column_names[case_study]["irb_consideration"]]
                key_factors = row[column_names[case_study]["key_factors"]].replace("-99", "")
                ethical_concerns = row[column_names[case_study]["ethical_concerns"]].replace("-99", "")
                response = CaseStudyResponse(respondent, irb_consideration, key_factors, ethical_concerns)
                if response.key_factors or response.ethical_concerns:
                    responses[case_study].append(response)

    document = docx.Document()

    document.add_heading("Prim&R 18 Survey Responses", 0)
    
    for i in range(1, 12):

        case_name = case_study_info[f"case_{i}"]["title"]
        case_prompt = case_study_info[f"case_{i}"]["prompt"]

        document.add_heading(f"Case {i}: {case_name}", 2)
        p = document.add_paragraph()
        p.add_run(case_prompt).italic = True
        p.add_run("\n")
        add_horizontal_line(p)

        for response in responses[f"case_{i}"]:
            document.add_heading(f"Case {i}: {case_name}", 4)
            document.add_heading("Respondent", 5)
            document.add_paragraph(response.respondent_id)
            document.add_heading("IRB Consideration", 5)
            document.add_paragraph(response.irb_consideration)
            document.add_heading("Key Factors", 5)
            document.add_paragraph(response.key_factors or "None")
            document.add_heading("Ethical concerns", 5)
            p = document.add_paragraph(response.ethical_concerns or "None")
            p.add_run("\n")
            add_horizontal_line(p)

        document.add_page_break()
        

    document.save(args.outfile)



if __name__ == "__main__":
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        '-i', '--infile',
        type=pathlib.Path,
        dest='infile',
        required=True,
        help='Location of the input CSV file containing survey data'
    )
    parser.add_argument(
        '-o', '--outfile',
        type=pathlib.Path,
        dest='outfile',
        default=pathlib.Path.cwd() / f'primr18_survey_text.docx',
        help='The location where report of survey free text responses will be saved'
    )
    args = parser.parse_args()
    main(args)
